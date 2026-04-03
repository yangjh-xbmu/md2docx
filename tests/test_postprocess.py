"""Tests for post-processing engine."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from md2docx.postprocess import _resolve_var, _apply_heading_numbering


class TestResolveVar:
    def test_basic_replacement(self) -> None:
        result = _resolve_var("{title} by {author}", {"title": "Hello", "author": "World"})
        assert result == "Hello by World"

    def test_no_placeholders(self) -> None:
        result = _resolve_var("plain text", {"title": "ignored"})
        assert result == "plain text"

    def test_missing_key(self) -> None:
        result = _resolve_var("{title} and {missing}", {"title": "Hello"})
        assert result == "Hello and {missing}"


class TestApplyHeadingNumbering:
    @pytest.mark.integration
    def test_numbering_applied(self, tmp_path: Path) -> None:
        """Create a minimal docx with headings and verify numbering."""
        doc = Document()
        doc.add_heading("Chapter One", level=1)
        doc.add_paragraph("Content")
        doc.add_heading("Section A", level=2)
        doc.add_heading("Section B", level=2)
        doc.add_heading("Chapter Two", level=1)

        _apply_heading_numbering(doc)

        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert headings[0].text == "1 Chapter One"
        assert headings[1].text == "1.1 Section A"
        assert headings[2].text == "1.2 Section B"
        assert headings[3].text == "2 Chapter Two"
