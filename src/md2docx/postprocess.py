"""Post-processing engine: python-docx based refinements on pandoc output."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from md2docx.style import PostProcessConfig


def postprocess(
    docx_path: Path,
    config: PostProcessConfig,
    metadata: dict[str, Any],
    output_path: Path | None = None,
) -> Path:
    """Apply all enabled post-processors to a docx file.

    Processors run in order: cover → toc → heading_numbering → image_resize → header_footer.
    """
    doc = Document(str(docx_path))
    out = output_path or docx_path

    if config.cover.enabled:
        _apply_cover(doc, config, metadata)

    if config.toc:
        _insert_toc(doc, config.toc_depth)

    if config.heading_numbering:
        _apply_heading_numbering(doc)

    if config.image_width_pct > 0:
        _resize_images(doc, config.image_width_pct)

    if config.header_left or config.header_right or config.footer_center:
        _apply_header_footer(doc, config, metadata)

    doc.save(str(out))
    return out


def _resolve_var(template: str, metadata: dict[str, Any]) -> str:
    """Replace {key} placeholders with metadata values."""
    result = template
    for key, value in metadata.items():
        result = result.replace(f"{{{key}}}", str(value))
    return result


def _apply_cover(
    doc: Document,
    config: PostProcessConfig,
    metadata: dict[str, Any],
) -> None:
    """Insert a cover page at the beginning of the document."""
    # Insert a section break before the first paragraph to create a cover page
    if not doc.paragraphs:
        return

    cover_fields = config.cover.fields
    resolved = {k: _resolve_var(v, metadata) for k, v in cover_fields.items()}

    # Insert cover content before the first paragraph
    first_para = doc.paragraphs[0]
    ref_element = first_para._element

    # Title
    title_text = resolved.get("title", metadata.get("title", ""))
    if title_text:
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)

        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "4800")
        pPr.append(spacing)

        p.append(pPr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "56")  # 28pt
        rPr.append(sz)
        b = OxmlElement("w:b")
        rPr.append(b)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = title_text
        r.append(t)
        p.append(r)
        ref_element.addprevious(p)

    # Author
    author_text = resolved.get("author", metadata.get("author", ""))
    if author_text:
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "600")
        pPr.append(spacing)
        p.append(pPr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "32")  # 16pt
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = author_text
        r.append(t)
        p.append(r)
        ref_element.addprevious(p)

    # Date
    date_text = resolved.get("date", metadata.get("date", ""))
    if date_text:
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "300")
        pPr.append(spacing)
        p.append(pPr)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = str(date_text)
        r.append(t)
        p.append(r)
        ref_element.addprevious(p)

    # Page break after cover
    p_break = OxmlElement("w:p")
    r_break = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r_break.append(br)
    p_break.append(r_break)
    ref_element.addprevious(p_break)


def _insert_toc(doc: Document, depth: int = 3) -> None:
    """Insert a TOC field code at the beginning of the document (after cover if any)."""
    # Find insertion point: after any page break, or at the start
    insert_before = doc.paragraphs[0]._element if doc.paragraphs else None
    if insert_before is None:
        return

    # TOC title paragraph
    toc_title = OxmlElement("w:p")
    toc_title_pPr = OxmlElement("w:pPr")
    toc_title_pStyle = OxmlElement("w:pStyle")
    toc_title_pStyle.set(qn("w:val"), "TOCHeading")
    toc_title_pPr.append(toc_title_pStyle)
    toc_title.append(toc_title_pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "目录"
    r.append(t)
    toc_title.append(r)

    # TOC field
    toc_para = OxmlElement("w:p")
    r_begin = OxmlElement("w:r")
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fldChar_begin)
    toc_para.append(r_begin)

    r_instr = OxmlElement("w:r")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f' TOC \\o "1-{depth}" \\h \\z \\u '
    r_instr.append(instrText)
    toc_para.append(r_instr)

    r_end = OxmlElement("w:r")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    r_end.append(fldChar_end)
    toc_para.append(r_end)

    # Page break after TOC
    p_break = OxmlElement("w:p")
    r_break = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r_break.append(br)
    p_break.append(r_break)

    insert_before.addprevious(toc_title)
    insert_before.addprevious(toc_para)
    insert_before.addprevious(p_break)


def _apply_heading_numbering(doc: Document) -> None:
    """Add hierarchical numbering to headings (1, 1.1, 1.1.1, etc.)."""
    counters = [0] * 9  # h1..h9

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if not style_name.startswith("Heading"):
            continue

        try:
            level = int(style_name.replace("Heading ", "").replace("Heading", "").strip())
        except ValueError:
            continue

        if level < 1 or level > 9:
            continue

        idx = level - 1
        counters[idx] += 1
        # Reset lower-level counters
        for i in range(idx + 1, 9):
            counters[i] = 0

        prefix = ".".join(str(counters[i]) for i in range(idx + 1))
        para.text = f"{prefix} {para.text}"


def _resize_images(doc: Document, width_pct: int) -> None:
    """Resize all inline images to a percentage of page width."""
    if width_pct <= 0 or width_pct > 100:
        return

    for section in doc.sections:
        page_width = section.page_width - section.left_margin - section.right_margin
        target_width = int(page_width * width_pct / 100)

        for shape in doc.inline_shapes:
            if shape.width and shape.width > target_width:
                ratio = target_width / shape.width
                shape.width = target_width
                shape.height = int(shape.height * ratio)
        break  # Use first section's dimensions


def _apply_header_footer(
    doc: Document,
    config: PostProcessConfig,
    metadata: dict[str, Any],
) -> None:
    """Set header and footer for all sections."""
    for section in doc.sections:
        # Header
        if config.header_left or config.header_right:
            header = section.header
            header.is_linked_to_previous = False

            if header.paragraphs:
                hp = header.paragraphs[0]
            else:
                hp = header.add_paragraph()

            hp.clear()

            left_text = _resolve_var(config.header_left, metadata) if config.header_left else ""
            right_text = _resolve_var(config.header_right, metadata) if config.header_right else ""

            if left_text and right_text:
                # Left-aligned text + tab + right-aligned text
                hp.alignment = None
                run_left = hp.add_run(left_text)
                run_left.font.size = Pt(9)
                run_tab = hp.add_run("\t\t")
                run_right = hp.add_run(right_text)
                run_right.font.size = Pt(9)
            elif left_text:
                hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = hp.add_run(left_text)
                run.font.size = Pt(9)
            elif right_text:
                hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = hp.add_run(right_text)
                run.font.size = Pt(9)

        # Footer
        if config.footer_center:
            footer = section.footer
            footer.is_linked_to_previous = False

            if footer.paragraphs:
                fp = footer.paragraphs[0]
            else:
                fp = footer.add_paragraph()

            fp.clear()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            footer_text = _resolve_var(config.footer_center, metadata)

            if "{page}" in footer_text:
                parts = footer_text.split("{page}")
                if parts[0]:
                    run = fp.add_run(parts[0])
                    run.font.size = Pt(9)

                # PAGE field
                r = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "18")
                rPr.append(sz)
                r.append(rPr)
                fldChar_begin = OxmlElement("w:fldChar")
                fldChar_begin.set(qn("w:fldCharType"), "begin")
                r.append(fldChar_begin)
                fp._element.append(r)

                r2 = OxmlElement("w:r")
                instrText = OxmlElement("w:instrText")
                instrText.text = " PAGE "
                r2.append(instrText)
                fp._element.append(r2)

                r3 = OxmlElement("w:r")
                fldChar_end = OxmlElement("w:fldChar")
                fldChar_end.set(qn("w:fldCharType"), "end")
                r3.append(fldChar_end)
                fp._element.append(r3)

                if len(parts) > 1 and parts[1]:
                    run = fp.add_run(parts[1])
                    run.font.size = Pt(9)
            else:
                run = fp.add_run(footer_text)
                run.font.size = Pt(9)
