"""Tests for CSS to Word style conversion."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, RGBColor

from md2docx.css2style import (
    build_reference_docx,
    css_to_reference_docx,
    parse_css,
    _parse_color,
    _length_to_pt,
    _is_cjk_font,
)


class TestParseCSS:
    def test_parse_heading(self) -> None:
        css = "h1 { font-size: 22pt; font-weight: bold; text-align: center; }"
        styles = parse_css(css)
        assert len(styles) == 1
        s = styles[0]
        assert s.word_style_name == "Heading 1"
        assert s.font_size == Pt(22)
        assert s.bold is True

    def test_parse_paragraph(self) -> None:
        css = "p { font-size: 12pt; text-indent: 2em; line-height: 1.75; }"
        styles = parse_css(css)
        assert len(styles) == 1
        s = styles[0]
        assert s.word_style_name == "Normal"
        assert s.font_size == Pt(12)
        assert s.line_height == 1.75

    def test_parse_class_selector(self) -> None:
        css = ".abstract { font-size: 10.5pt; margin-left: 2cm; }"
        styles = parse_css(css)
        assert len(styles) == 1
        assert styles[0].word_style_name == "abstract"

    def test_parse_font_family_cjk(self) -> None:
        css = 'p { font-family: "SimSun", "Times New Roman"; }'
        styles = parse_css(css)
        s = styles[0]
        assert s.font_family == "Times New Roman"
        assert s.font_family_eastasia == "SimSun"

    def test_parse_color(self) -> None:
        css = "p { color: #333333; }"
        styles = parse_css(css)
        assert styles[0].color == RGBColor(0x33, 0x33, 0x33)

    def test_parse_multiple_rules(self) -> None:
        css = "h1 { font-size: 22pt; } h2 { font-size: 16pt; } p { font-size: 12pt; }"
        styles = parse_css(css)
        assert len(styles) == 3


class TestParseColor:
    def test_hex6(self) -> None:
        import tinycss2

        tokens = tinycss2.parse_component_value_list("#ff0000")
        assert _parse_color(tokens) == RGBColor(255, 0, 0)

    def test_hex3(self) -> None:
        import tinycss2

        tokens = tinycss2.parse_component_value_list("#f00")
        assert _parse_color(tokens) == RGBColor(255, 0, 0)

    def test_named(self) -> None:
        import tinycss2

        tokens = tinycss2.parse_component_value_list("blue")
        assert _parse_color(tokens) == RGBColor(0, 0, 255)


class TestLengthConversion:
    def test_pt(self) -> None:
        assert _length_to_pt(12.0, "pt") == 12.0

    def test_cm(self) -> None:
        result = _length_to_pt(1.0, "cm")
        assert abs(result - 28.3465) < 0.01

    def test_inch(self) -> None:
        assert _length_to_pt(1.0, "in") == 72.0


class TestIsCJKFont:
    def test_simsun(self) -> None:
        assert _is_cjk_font("SimSun") is True

    def test_times(self) -> None:
        assert _is_cjk_font("Times New Roman") is False

    def test_chinese_name(self) -> None:
        assert _is_cjk_font("宋体") is True


class TestBuildReferenceDocx:
    def test_build_from_parsed(self, tmp_path: Path) -> None:
        css = "h1 { font-size: 22pt; font-weight: bold; } p { font-size: 12pt; }"
        styles = parse_css(css)
        output = tmp_path / "ref.docx"
        result = build_reference_docx(styles, output=output)
        assert result.exists()

        doc = Document(str(result))
        h1_style = doc.styles["Heading 1"]
        assert h1_style.font.size == Pt(22)
        assert h1_style.font.bold is True

    def test_build_with_cjk_font(self, tmp_path: Path) -> None:
        css = 'p { font-family: "SimSun", "Times New Roman"; font-size: 12pt; }'
        styles = parse_css(css)
        output = tmp_path / "ref.docx"
        build_reference_docx(styles, output=output)
        doc = Document(str(output))
        normal = doc.styles["Normal"]
        assert normal.font.name == "Times New Roman"


class TestCssToReferenceDocx:
    def test_from_fixture(self, fixtures_dir: Path, tmp_path: Path) -> None:
        css_file = fixtures_dir / "academic.css"
        output = tmp_path / "academic_ref.docx"
        result = css_to_reference_docx(css_file, output=output)
        assert result.exists()

        doc = Document(str(result))
        style_names = [s.name for s in doc.styles]
        assert "Heading 1" in style_names
        assert "abstract" in style_names
        assert "keywords" in style_names

    def test_empty_css_raises(self, tmp_path: Path) -> None:
        css_file = tmp_path / "empty.css"
        css_file.write_text("/* no rules */", encoding="utf-8")
        with pytest.raises(ValueError, match="没有找到"):
            css_to_reference_docx(css_file)
