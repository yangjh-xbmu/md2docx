"""Integration tests for the full conversion pipeline."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document


@pytest.mark.integration
class TestConvert:
    def test_simple_convert(self, simple_md: Path, tmp_output: Path) -> None:
        from md2docx.convert import convert

        result = convert(source=simple_md, output=tmp_output)
        assert result.exists()
        assert result.suffix == ".docx"

        doc = Document(str(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Hello World" in full_text
        assert "Section One" in full_text

    def test_convert_with_frontmatter(self, frontmatter_md: Path, tmp_output: Path) -> None:
        from md2docx.convert import convert

        result = convert(source=frontmatter_md, output=tmp_output)
        assert result.exists()

        doc = Document(str(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "概述" in full_text

    def test_convert_no_post(self, simple_md: Path, tmp_output: Path) -> None:
        from md2docx.convert import convert

        result = convert(source=simple_md, output=tmp_output, no_post=True)
        assert result.exists()

    def test_convert_default_output(self, simple_md: Path, tmp_path: Path) -> None:
        """When no output specified, should create .docx next to source."""
        from md2docx.convert import convert

        # Copy simple.md to tmp to avoid polluting fixtures
        src = tmp_path / "test.md"
        src.write_text(simple_md.read_text(encoding="utf-8"), encoding="utf-8")

        result = convert(source=src)
        assert result == tmp_path / "test.docx"
        assert result.exists()


@pytest.mark.integration
class TestConvertMerged:
    def test_merge_from_contents_yaml(self, merge_dir: Path, tmp_output: Path) -> None:
        from md2docx.convert import convert_merged

        result = convert_merged(
            contents=merge_dir / "contents.yaml",
            output=tmp_output,
        )
        assert result.exists()

        doc = Document(str(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "第一章 引言" in full_text
        assert "第二章 方法" in full_text

    def test_merge_from_glob(self, merge_dir: Path, tmp_output: Path) -> None:
        from md2docx.convert import convert_merged

        result = convert_merged(
            patterns=[str(merge_dir / "chapter*.md")],
            output=tmp_output,
        )
        assert result.exists()
